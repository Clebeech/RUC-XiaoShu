from __future__ import annotations

import html
import json
import mimetypes
import re
from pathlib import Path

from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile, status
from pypdf import PdfReader
from sqlalchemy.orm import Session

from ..config import settings
from ..models import Document, DocumentChunk, KnowledgeBase
from .embedding_service import EmbeddingService


class DocumentService:
    chunk_size = 1000
    overlap = 200

    @classmethod
    async def ingest_upload(cls, upload: UploadFile, knowledge_base: KnowledgeBase) -> tuple[Document, int]:
        payload = await upload.read()
        await upload.seek(0)
        file_name = upload.filename or "upload.txt"
        content = cls._extract_text_from_payload(
            payload=payload,
            file_name=file_name,
            content_type=upload.content_type,
        )
        if not content.strip():
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Document is empty after parsing")

        upload_dir = Path(settings.upload_dir)
        upload_dir.mkdir(parents=True, exist_ok=True)
        target_path = cls._build_storage_path(upload_dir, file_name)
        target_path.write_bytes(payload)

        document = Document(
            knowledge_base=knowledge_base,
            name=file_name,
            content=content,
            file_path=str(target_path),
            size=len(payload),
            mime_type=upload.content_type or "application/octet-stream",
        )

        chunk_payloads = await cls.build_chunk_payloads(content, file_name)
        for chunk in chunk_payloads:
            document.chunks.append(
                DocumentChunk(
                    content=chunk["content"],
                    embedding_json=json.dumps(chunk["embedding"]),
                    section=chunk["section"],
                    start_index=chunk["start_index"],
                    end_index=chunk["end_index"],
                )
            )

        return document, len(content)

    @classmethod
    async def ingest_local_file(
        cls,
        file_path: Path,
        knowledge_base: KnowledgeBase,
        document_name: str | None = None,
    ) -> tuple[Document, int]:
        payload = file_path.read_bytes()
        mime_type = mimetypes.guess_type(file_path.name)[0] or "application/octet-stream"
        content = cls._extract_text_from_payload(
            payload=payload,
            file_name=file_path.name,
            content_type=mime_type,
        )
        if not content.strip():
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Document is empty after parsing")

        upload_dir = Path(settings.upload_dir)
        upload_dir.mkdir(parents=True, exist_ok=True)
        target_name = document_name or file_path.name
        target_path = cls._build_storage_path(upload_dir, file_path.name)
        target_path.write_bytes(payload)

        document = Document(
            knowledge_base=knowledge_base,
            name=target_name,
            content=content,
            file_path=str(target_path),
            size=len(payload),
            mime_type=mime_type,
        )

        chunk_payloads = await cls.build_chunk_payloads(content, target_name)
        for chunk in chunk_payloads:
            document.chunks.append(
                DocumentChunk(
                    content=chunk["content"],
                    embedding_json=json.dumps(chunk["embedding"]),
                    section=chunk["section"],
                    start_index=chunk["start_index"],
                    end_index=chunk["end_index"],
                )
            )

        return document, len(content)

    @classmethod
    async def reindex_knowledge_base(cls, db: Session, knowledge_base: KnowledgeBase) -> int:
        count = 0
        for document in knowledge_base.documents:
            chunk_payloads = await cls.build_chunk_payloads(document.content, document.name)
            document.chunks.clear()
            for chunk in chunk_payloads:
                document.chunks.append(
                    DocumentChunk(
                        content=chunk["content"],
                        embedding_json=json.dumps(chunk["embedding"]),
                        section=chunk["section"],
                        start_index=chunk["start_index"],
                        end_index=chunk["end_index"],
                    )
                )
            count += 1

        db.flush()
        return count

    @classmethod
    async def _extract_text(cls, upload: UploadFile) -> str:
        payload = await upload.read()
        await upload.seek(0)
        return cls._extract_text_from_payload(
            payload=payload,
            file_name=upload.filename or "",
            content_type=upload.content_type,
        )

    @classmethod
    def _extract_text_from_payload(cls, payload: bytes, file_name: str, content_type: str | None = None) -> str:
        content_type = (content_type or "").lower()
        suffix = Path(file_name).suffix.lower()
        if content_type.startswith("text/") or suffix in {".txt", ".md"}:
            return payload.decode("utf-8", errors="ignore")
        if suffix == ".pdf":
            return cls._read_pdf(payload)
        if suffix in {".docx", ".doc"}:
            return cls._read_docx(payload)

        return payload.decode("utf-8", errors="ignore")

    @staticmethod
    def _read_pdf(payload: bytes) -> str:
        from io import BytesIO

        reader = PdfReader(BytesIO(payload))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(page.strip() for page in pages if page.strip())

    @staticmethod
    def _read_docx(payload: bytes) -> str:
        from io import BytesIO

        document = DocxDocument(BytesIO(payload))
        return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())

    @classmethod
    def _split_into_chunks(cls, content: str, file_name: str) -> list[dict[str, str | int]]:
        paragraphs = cls._normalize_paragraphs(content)
        chunks: list[dict[str, str | int]] = []
        current = ""
        start_index = 0

        for paragraph in paragraphs:
            candidate = f"{current}\n\n{paragraph}".strip() if current else paragraph
            if len(candidate) <= cls.chunk_size:
                current = candidate
                continue

            if current:
                chunks.append(cls._make_chunk(current, file_name, len(chunks), start_index))
                overlap_text = current[-cls.overlap :] if len(current) > cls.overlap else current
                start_index = max(0, start_index + len(current) - len(overlap_text))
                current = f"{overlap_text}\n\n{paragraph}".strip()
            else:
                chunks.append(cls._make_chunk(paragraph[: cls.chunk_size], file_name, len(chunks), start_index))
                start_index += cls.chunk_size
                current = paragraph[cls.chunk_size - cls.overlap :].strip()

        if current:
            chunks.append(cls._make_chunk(current, file_name, len(chunks), start_index))

        return chunks

    @classmethod
    def _normalize_paragraphs(cls, content: str) -> list[str]:
        raw_paragraphs = [paragraph.strip() for paragraph in content.split("\n\n") if paragraph.strip()]
        normalized: list[str] = []
        max_unit = max(1, cls.chunk_size - cls.overlap)

        for paragraph in raw_paragraphs:
            if len(paragraph) <= cls.chunk_size:
                normalized.append(paragraph)
                continue

            normalized.extend(cls._split_large_paragraph(paragraph, max_unit))

        return normalized

    @classmethod
    def _split_large_paragraph(cls, paragraph: str, max_unit: int) -> list[str]:
        line_based = [line.strip() for line in paragraph.splitlines() if line.strip()]
        if len(line_based) > 1:
            parts: list[str] = []
            current = ""
            for line in line_based:
                candidate = f"{current}\n{line}".strip() if current else line
                if len(candidate) <= max_unit:
                    current = candidate
                    continue
                if current:
                    parts.append(current)
                if len(line) <= max_unit:
                    current = line
                else:
                    parts.extend(cls._split_by_length(line, max_unit))
                    current = ""
            if current:
                parts.append(current)
            return parts

        return cls._split_by_length(paragraph, max_unit)

    @staticmethod
    def _split_by_length(text: str, max_unit: int) -> list[str]:
        parts: list[str] = []
        cursor = 0
        while cursor < len(text):
            parts.append(text[cursor : cursor + max_unit].strip())
            cursor += max_unit
        return [part for part in parts if part]

    @classmethod
    async def build_chunk_payloads(cls, content: str, file_name: str) -> list[dict[str, str | int | list[float]]]:
        normalized_content = cls._normalize_content_for_indexing(content)
        chunks = cls._split_into_chunks(normalized_content, file_name)
        embeddings = await EmbeddingService.embed_texts(
            [str(chunk["content"]) for chunk in chunks],
            text_type="document",
        )

        return [
            {
                **chunk,
                "embedding": embedding,
            }
            for chunk, embedding in zip(chunks, embeddings, strict=False)
        ]

    @staticmethod
    def _make_chunk(content: str, file_name: str, index: int, start_index: int) -> dict[str, str | int]:
        clean = content.strip()
        return {
            "content": clean,
            "section": f"{file_name} - 第{index + 1}部分",
            "start_index": start_index,
            "end_index": start_index + len(clean),
        }

    @staticmethod
    def _build_storage_path(upload_dir: Path, file_name: str) -> Path:
        candidate = upload_dir / Path(file_name).name
        if not candidate.exists():
            return candidate

        stem = candidate.stem
        suffix = candidate.suffix
        index = 1
        while True:
            attempt = upload_dir / f"{stem}_{index}{suffix}"
            if not attempt.exists():
                return attempt
            index += 1

    @classmethod
    def _normalize_content_for_indexing(cls, content: str) -> str:
        normalized = content.replace("\r\n", "\n").replace("\r", "\n")
        normalized = re.sub(r"!\[[^\]]*\]\([^)]+\)", " ", normalized)
        normalized = re.sub(r"<table\b.*?</table>", cls._table_to_text, normalized, flags=re.IGNORECASE | re.DOTALL)
        normalized = re.sub(r"<br\s*/?>", "\n", normalized, flags=re.IGNORECASE)
        normalized = re.sub(r"</p\s*>", "\n", normalized, flags=re.IGNORECASE)
        normalized = re.sub(r"<[^>]+>", " ", normalized)
        normalized = html.unescape(normalized)
        normalized = re.sub(r"[ \t]+", " ", normalized)
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        return normalized.strip()

    @classmethod
    def _table_to_text(cls, match: re.Match[str]) -> str:
        table_html = match.group(0)
        rows = re.findall(r"<tr\b[^>]*>(.*?)</tr>", table_html, flags=re.IGNORECASE | re.DOTALL)
        formatted_rows: list[str] = []
        for row in rows:
            cells = re.findall(r"<t[dh]\b[^>]*>(.*?)</t[dh]>", row, flags=re.IGNORECASE | re.DOTALL)
            clean_cells = [cls._clean_html_cell(cell) for cell in cells]
            clean_cells = [cell for cell in clean_cells if cell]
            if clean_cells:
                formatted_rows.append(" | ".join(clean_cells))
        if not formatted_rows:
            return "\n"
        return "\n" + "\n".join(formatted_rows) + "\n"

    @staticmethod
    def _clean_html_cell(value: str) -> str:
        text = re.sub(r"<br\s*/?>", " / ", value, flags=re.IGNORECASE)
        text = re.sub(r"<[^>]+>", " ", text)
        text = html.unescape(text)
        text = re.sub(r"\s+", " ", text).strip()
        return text
